# =====================================================
# 모듈 3 - 주석 초안 자동 생성 프로그램
# 분기 결산명세서의 숫자를 읽어서 주석 초안 문구를 자동으로 만듭니다.
# 판단이 필요한 부분은 [작성자 확인 필요] 표시를 붙입니다.
# =====================================================

import pandas as pd                          # 표 데이터를 다루는 도구
from openpyxl.styles import PatternFill, Font, Alignment  # 엑셀 스타일 도구
from docx import Document                    # 워드 파일 만드는 도구
from docx.shared import Pt, RGBColor        # 글씨 크기, 색상 도구
import os                                     # 파일/폴더 경로 처리 도구

# -------------------------------------------------------
# [설정] 여기서 원하는 값을 바꾸세요!
# -------------------------------------------------------

# 입력: module2에서 만든 분기 결산명세서 파일
STATEMENT_FILE = "data/output/분기결산명세서_2026년1분기.xlsx"

# 출력 파일 경로
OUTPUT_WORD  = "data/output/주석_초안_2026년1분기.docx"   # 워드 파일
OUTPUT_EXCEL = "data/output/주석_초안_2026년1분기.xlsx"   # 엑셀 파일

# 문서에 들어갈 기본 정보
COMPANY_NAME   = "주식회사 OO"          # 회사명을 입력하세요
REPORT_DATE    = "2026년 3월 31일"     # 보고 기준일
QUARTER_LABEL  = "2026년 1분기"        # 당기 표시
PRIOR_LABEL    = "2025년 1분기"        # 전기 표시

# 중요성 기준: 이 금액 이상인 항목은 중요 항목으로 표시
MATERIAL_THRESHOLD = 100_000_000  # 1억 원 (필요시 수정)

# 변동률 기준: 이 비율 이상 변동하면 변동 사유 기재 필요 표시
VARIANCE_THRESHOLD = 0.20  # 20%

# -------------------------------------------------------
# [계정별 표준 문구 템플릿]
# {당기금액}, {전기금액}, {증감액}, {증감률}, {방향} 자리에 숫자가 자동으로 들어갑니다.
# -------------------------------------------------------
PHRASE_TEMPLATES = {
    "4010": "제품 매출액은 {당기금액}원으로, 전기({전기금액}원) 대비 {증감액}원({증감률}) {방향}하였습니다.",
    "4020": "상품 매출액은 {당기금액}원으로, 전기({전기금액}원) 대비 {증감액}원({증감률}) {방향}하였습니다.",
    "4030": "용역 매출액은 {당기금액}원으로, 전기({전기금액}원) 대비 {증감액}원({증감률}) {방향}하였습니다.",
    "4110": "기타 영업수익은 {당기금액}원으로, 전기({전기금액}원) 대비 {증감액}원({증감률}) {방향}하였습니다.",
    "4120": "임대 수익은 {당기금액}원으로, 전기({전기금액}원) 대비 {증감액}원({증감률}) {방향}하였습니다.",
    # 위에 없는 계정코드는 아래 기본 문구가 사용됩니다
    "DEFAULT": "{계정명}은(는) {당기금액}원으로, 전기({전기금액}원) 대비 {증감액}원({증감률}) {방향}하였습니다.",
}


# -------------------------------------------------------
# [함수] 결산명세서를 불러옵니다
# -------------------------------------------------------
def load_statement(file_path):
    """
    module2에서 만든 분기 결산명세서 엑셀 파일을 불러옵니다.
    입력: 파일 경로
    출력: pandas DataFrame
    """
    print(f"  📂 결산명세서 불러오는 중... ({file_path})")

    if not os.path.exists(file_path):
        print(f"  [오류] 파일을 찾을 수 없습니다: {file_path}")
        return None

    df = pd.read_excel(file_path, sheet_name="결산명세서")

    # 필요한 열이 있는지 확인
    required_cols = ["계정코드", "계정명", "당기금액"]
    for col in required_cols:
        if col not in df.columns:
            print(f"  [오류] '{col}' 열을 찾을 수 없습니다.")
            return None

    # 전기금액 열이 없으면 0으로 채우기
    if "전기금액" not in df.columns:
        df["전기금액"] = 0
    if "증감액" not in df.columns:
        df["증감액"] = df["당기금액"] - df["전기금액"]
    if "증감률" not in df.columns:
        df["증감률"] = df.apply(
            lambda row: row["증감액"] / row["전기금액"] if row["전기금액"] != 0 else None,
            axis=1
        )

    print(f"  [완료] {len(df)}개 계정 불러옴")
    return df


# -------------------------------------------------------
# [함수] 숫자를 읽기 쉬운 형태로 변환합니다
# -------------------------------------------------------
def format_amount(amount):
    """
    숫자를 콤마가 있는 형태로 변환합니다.
    예: 1327500000 → "1,327,500,000"
    입력: 숫자
    출력: 문자열
    """
    if amount is None or (isinstance(amount, float) and pd.isna(amount)):
        return "0"
    return f"{int(amount):,}"


# -------------------------------------------------------
# [함수] 계정별 표준 문구를 생성합니다
# -------------------------------------------------------
def build_standard_phrase(account_code, account_name, current, prior, variance, variance_rate):
    """
    계정 정보를 받아서 주석 문구를 자동으로 만듭니다.
    입력: 계정코드, 계정명, 당기금액, 전기금액, 증감액, 증감률
    출력: (자동 문구 문자열, 판단 필요 여부 bool)
    """
    # 증감 방향 (증가/감소)
    if variance > 0:
        direction = "증가"
    elif variance < 0:
        direction = "감소"
    else:
        direction = "변동 없음"

    # 증감률 텍스트 (없으면 "-")
    if variance_rate is not None and not pd.isna(variance_rate):
        rate_text = f"{abs(variance_rate):.1%}"
    else:
        rate_text = "-"

    # 템플릿 선택 (계정코드에 맞는 것 또는 기본 템플릿)
    template = PHRASE_TEMPLATES.get(str(account_code), PHRASE_TEMPLATES["DEFAULT"])

    # 숫자를 문구에 끼워넣기
    phrase = template.format(
        계정명=account_name,
        당기금액=format_amount(current),
        전기금액=format_amount(prior),
        증감액=format_amount(abs(variance)),
        증감률=rate_text,
        방향=direction,
    )

    # 판단이 필요한지 여부 결정
    needs_review = False
    review_reason = []

    if abs(current) >= MATERIAL_THRESHOLD:  # 중요성 기준 초과
        needs_review = True
        review_reason.append(f"금액이 {format_amount(MATERIAL_THRESHOLD)}원 이상인 중요 항목")

    if variance_rate is not None and not pd.isna(variance_rate) and abs(variance_rate) > VARIANCE_THRESHOLD:
        needs_review = True
        review_reason.append(f"전기 대비 {abs(variance_rate):.1%} 변동")

    return phrase, needs_review, review_reason


# -------------------------------------------------------
# [함수] 워드 파일로 주석 초안을 만듭니다
# -------------------------------------------------------
def write_word_draft(df, output_path):
    """
    분석 결과를 워드(.docx) 파일로 저장합니다.
    입력: 결산명세서 DataFrame, 저장 경로
    출력: 없음 (파일 저장)
    """
    print(f"  📝 워드 초안 작성 중... ({output_path})")

    doc = Document()  # 새 워드 문서 만들기

    # === 제목 ===
    title = doc.add_heading(f"{COMPANY_NAME}", level=1)
    title.alignment = 1  # 가운데 정렬
    subtitle = doc.add_heading(f"{QUARTER_LABEL} 주석 초안 (자동 생성)", level=2)
    subtitle.alignment = 1

    doc.add_paragraph(f"기준일: {REPORT_DATE}")
    doc.add_paragraph(f"※ 이 문서는 자동으로 생성된 초안입니다. [작성자 확인 필요] 항목을 반드시 검토하세요.")
    doc.add_paragraph("")  # 빈 줄

    needs_review_count = 0  # 확인 필요 항목 수

    # === 각 계정별 주석 문구 ===
    for _, row in df.iterrows():
        code = str(row["계정코드"])
        name = str(row["계정명"])
        current = row["당기금액"]
        prior = row.get("전기금액", 0)
        variance = row.get("증감액", current - prior)
        variance_rate = row.get("증감률", None)

        # 표준 문구 생성
        phrase, needs_review, review_reasons = build_standard_phrase(
            code, name, current, prior, variance, variance_rate
        )

        # 계정 제목 (굵게)
        heading = doc.add_heading(f"{name} ({code})", level=3)

        # 자동 생성 문구
        p = doc.add_paragraph(phrase)

        # 판단 필요 항목 빨간색으로 표시
        if needs_review:
            needs_review_count += 1
            reasons_text = ", ".join(review_reasons)
            review_para = doc.add_paragraph()
            run = review_para.add_run(f"[작성자 확인 필요: {reasons_text} — 변동 주요 원인을 기재하세요]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 빨간색
            run.font.bold = True

        doc.add_paragraph("")  # 항목 사이 빈 줄

    # === 마지막 요약 ===
    doc.add_heading("작성 완료 요약", level=2)
    doc.add_paragraph(f"• 총 계정 수: {len(df)}개")
    doc.add_paragraph(f"• 자동 문구 생성: {len(df)}개")
    doc.add_paragraph(f"• [작성자 확인 필요] 항목: {needs_review_count}개")

    # 저장
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"  [완료] 워드 파일 저장 완료!")
    return needs_review_count


# -------------------------------------------------------
# [함수] 엑셀 파일로도 저장합니다
# -------------------------------------------------------
def write_excel_draft(df, output_path):
    """
    주석 초안을 엑셀 파일로도 저장합니다.
    입력: 결산명세서 DataFrame, 저장 경로
    출력: 없음 (파일 저장)
    """
    print(f"  📊 엑셀 초안 저장 중... ({output_path})")

    rows = []  # 결과 행들을 담을 목록

    for _, row in df.iterrows():
        code = str(row["계정코드"])
        name = str(row["계정명"])
        current = row["당기금액"]
        prior = row.get("전기금액", 0)
        variance = row.get("증감액", current - prior)
        variance_rate = row.get("증감률", None)

        phrase, needs_review, review_reasons = build_standard_phrase(
            code, name, current, prior, variance, variance_rate
        )

        rows.append({
            "계정코드": code,
            "계정명": name,
            "당기금액": current,
            "전기금액": prior,
            "증감액": variance,
            "증감률": variance_rate,
            "자동생성문구": phrase,
            "확인필요여부": "✅ 확인 필요" if needs_review else "자동완성",
            "확인사항": ", ".join(review_reasons) if review_reasons else "-",
        })

    result_df = pd.DataFrame(rows)

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        result_df.to_excel(writer, sheet_name="주석초안", index=False)
        ws = writer.sheets["주석초안"]

        # 헤더 스타일
        header_fill = PatternFill("solid", fgColor="4472C4")
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal="center")

        # 확인 필요 행 노란색 표시
        yellow_fill = PatternFill("solid", fgColor="FFD700")
        for row in ws.iter_rows(min_row=2):
            if row[7].value == "✅ 확인 필요":  # H열 = 확인필요여부
                for cell in row:
                    cell.fill = yellow_fill

        # 숫자 서식
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                if cell.column_letter in ["C", "D", "E"]:
                    if isinstance(cell.value, (int, float)):
                        cell.number_format = "#,##0"
                elif cell.column_letter == "F":
                    if isinstance(cell.value, (int, float)):
                        cell.number_format = "0.0%"

        # 열 너비 자동 조정
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)

    print(f"  [완료] 엑셀 파일 저장 완료!")


# -------------------------------------------------------
# [메인] 프로그램 실행 시 이 부분이 동작합니다
# -------------------------------------------------------
def main():
    print("=" * 55)
    print("  주석 초안 자동 생성 프로그램 시작!")
    print("=" * 55)
    print(f"  대상: {QUARTER_LABEL}")
    print(f"  회사: {COMPANY_NAME}")
    print()

    # 1단계: 결산명세서 불러오기
    df = load_statement(STATEMENT_FILE)
    if df is None:
        return
    print()

    # 2단계: 워드 초안 생성
    review_count = write_word_draft(df, OUTPUT_WORD)
    print()

    # 3단계: 엑셀 초안 생성
    write_excel_draft(df, OUTPUT_EXCEL)

    # 최종 요약
    print()
    print("=" * 55)
    print(f"  총 {len(df)}개 계정 주석 초안 생성 완료")
    print(f"  워드 파일: {OUTPUT_WORD}")
    print(f"  엑셀 파일: {OUTPUT_EXCEL}")
    if review_count > 0:
        print(f"  ⚠️  [작성자 확인 필요] 항목 {review_count}건 → 빨간색 항목 검토 필요")
    print("=" * 55)


# 프로그램 시작점
if __name__ == "__main__":
    main()
